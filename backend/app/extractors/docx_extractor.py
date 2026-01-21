"""DOCX text extraction using python-docx."""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import Optional
from pathlib import Path
from io import BytesIO


class DOCXExtractor:
    """Extract text content from DOCX files."""

    def __init__(self):
        self.supported_extensions = [".docx"]

    def extract(self, file_path: str | Path) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.

        Raises:
            FileNotFoundError: If file doesn't exist.
            ValueError: If file is not a valid DOCX.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.suffix.lower() not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        return self._extract_text(file_path)

    def extract_from_bytes(self, content: bytes, filename: str = "document.docx") -> str:
        """
        Extract text from DOCX bytes.

        Args:
            content: DOCX file content as bytes.
            filename: Original filename (for error messages).

        Returns:
            Extracted text content.
        """
        try:
            doc = Document(BytesIO(content))
            return self._process_document(doc)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX '{filename}': {e}")

    def _extract_text(self, file_path: Path) -> str:
        """Extract text from a DOCX file path."""
        try:
            doc = Document(file_path)
            return self._process_document(doc)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX '{file_path}': {e}")

    def _process_document(self, doc: Document) -> str:
        """Process a python-docx document and extract text."""
        text_parts = []

        for element in doc.element.body:
            if element.tag.endswith("p"):
                para = Paragraph(element, doc)
                text = self._extract_paragraph(para)
                if text:
                    text_parts.append(text)
            elif element.tag.endswith("tbl"):
                table = Table(element, doc)
                text = self._extract_table(table)
                if text:
                    text_parts.append(text)

        return "\n\n".join(text_parts)

    def _extract_paragraph(self, para: Paragraph) -> str:
        """Extract text from a paragraph with basic formatting."""
        if not para.text.strip():
            return ""

        # Check for heading style
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            level = style_name.replace("Heading ", "").strip()
            try:
                level_num = int(level)
                prefix = "#" * level_num
                return f"{prefix} {para.text}"
            except ValueError:
                pass

        # Check for list items
        if para._element.pPr is not None:
            numPr = para._element.pPr.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if numPr is not None:
                return f"- {para.text}"

        return para.text

    def _extract_table(self, table: Table) -> str:
        """Extract text from a table in markdown format."""
        rows = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            row_text = "| " + " | ".join(cells) + " |"
            rows.append(row_text)

            # Add header separator after first row
            if i == 0:
                separator = "|" + "|".join(["---" for _ in cells]) + "|"
                rows.append(separator)

        return "\n".join(rows)

    def get_metadata(self, file_path: str | Path) -> dict:
        """
        Extract metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Dictionary with metadata.
        """
        file_path = Path(file_path)
        doc = Document(file_path)
        core_props = doc.core_properties

        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
        }


# Singleton instance
docx_extractor = DOCXExtractor()
